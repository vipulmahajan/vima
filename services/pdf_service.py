"""Resume / report rendering — PDF (WeasyPrint) + DOCX (python-docx).

WeasyPrint is imported lazily inside each render call so that the app can
boot on machines that don't have the GTK runtime libraries installed
(common on developer Windows boxes). python-docx is pure-Python so it
imports eagerly.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any, Iterable

from jinja2 import Environment, FileSystemLoader, select_autoescape

# python-docx — pure Python, safe to import at module load.
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"

_env = Environment(
    loader=FileSystemLoader(str(TEMPLATES_DIR)),
    autoescape=select_autoescape(["html", "xml"]),
)


# ─────────────────────────────────────────────────────────────────────────────
# PDF (existing)
# ─────────────────────────────────────────────────────────────────────────────

import logging as _logging
_log = _logging.getLogger(__name__)

# GTK-related error strings that WeasyPrint surfaces when the runtime libraries
# are missing (common on Railway without the full GTK stack installed).
_GTK_ERROR_HINTS = (
    "cannot load library",
    "libgobject",
    "libpango",
    "libcairo",
    "gtk",
    "OSError",
    "no library called",
    "pangocairo",
)


def _is_gtk_error(exc: BaseException) -> bool:
    msg = str(exc).lower()
    return isinstance(exc, OSError) or any(h in msg for h in _GTK_ERROR_HINTS)


def render_resume_pdf(resume_data: dict[str, Any]) -> bytes:
    """Render a tailored resume to PDF bytes."""
    from weasyprint import HTML, CSS  # lazy: needs GTK runtime

    template = _env.get_template("resume.html")
    html_str = template.render(resume=resume_data)
    css_path = TEMPLATES_DIR / "resume.css"
    stylesheets = [CSS(filename=str(css_path))] if css_path.exists() else None
    return HTML(string=html_str, base_url=str(TEMPLATES_DIR)).write_pdf(
        stylesheets=stylesheets
    )


def try_render_resume_pdf(resume_data: dict[str, Any]) -> "bytes | None":
    """Like render_resume_pdf but returns None on GTK/OSError instead of raising."""
    try:
        return render_resume_pdf(resume_data)
    except Exception as exc:  # noqa: BLE001
        if _is_gtk_error(exc):
            _log.warning("PDF skipped — GTK runtime unavailable: %s", exc)
            return None
        raise


def render_interview_report_pdf(report_data: dict[str, Any]) -> bytes:
    """Render an end-of-interview feedback report to PDF bytes."""
    from weasyprint import HTML  # lazy: needs GTK runtime

    template = _env.get_template("interview_report.html")
    html_str = template.render(report=report_data)
    return HTML(string=html_str, base_url=str(TEMPLATES_DIR)).write_pdf()


def render_interview_prep_pdf(prep_data: dict[str, Any]) -> bytes:
    """Render an interview prep kit to PDF bytes."""
    from weasyprint import HTML, CSS  # lazy: needs GTK runtime

    template = _env.get_template("interview_prep.html")
    html_str = template.render(prep=prep_data)
    css_paths = [
        TEMPLATES_DIR / "resume.css",
        TEMPLATES_DIR / "interview_prep.css",
    ]
    stylesheets = [CSS(filename=str(p)) for p in css_paths if p.exists()]
    return HTML(string=html_str, base_url=str(TEMPLATES_DIR)).write_pdf(
        stylesheets=stylesheets or None
    )


def try_render_interview_prep_pdf(prep_data: dict[str, Any]) -> "bytes | None":
    """Like render_interview_prep_pdf but returns None on GTK/OSError instead of raising."""
    try:
        return render_interview_prep_pdf(prep_data)
    except Exception as exc:  # noqa: BLE001
        if _is_gtk_error(exc):
            _log.warning("Interview prep PDF skipped — GTK runtime unavailable: %s", exc)
            return None
        raise


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

# Theme palette for the "executive" template — kept subtle on purpose.
_FONT_NAME       = "Cambria"
_FONT_FALLBACK   = "Georgia"
_NAME_PT         = 24
_HEADLINE_PT     = 11
_CONTACT_PT      = 9.5
_SECTION_PT      = 11
_BODY_PT         = 10.5
_DATES_PT        = 9.5
_TITLE_COLOR     = RGBColor(0x1A, 0x1A, 0x1A)
_MUTED_COLOR     = RGBColor(0x55, 0x55, 0x55)
_DIVIDER_COLOR   = RGBColor(0x99, 0x99, 0x99)


def render_resume_docx(
    resume_data: dict[str, Any],
    template: str = "executive",
) -> bytes:
    """Render a resume as a clean, editable Word document.

    `template` is reserved for future variants; today only "executive" is
    implemented (clean serif, subtle dividers, no colour blocks).
    """
    if template != "executive":
        # Don't fail — just fall through. We can add more templates later.
        pass

    doc = Document()
    _configure_document(doc)

    r = resume_data or {}

    _add_header(doc, r)
    _add_summary(doc, r)
    _add_skills(doc, r)
    _add_experience(doc, r)
    _add_achievements(doc, r)
    _add_projects(doc, r)
    _add_education(doc, r)
    _add_certifications(doc, r)
    _add_publications(doc, r)
    _add_awards(doc, r)
    _add_memberships(doc, r)
    _add_languages(doc, r)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Document setup ──────────────────────────────────────────────────────────

def _configure_document(doc: "Document") -> None:
    """Set page margins and the default Normal style."""
    for section in doc.sections:
        section.top_margin    = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = _FONT_NAME
    style.font.size = Pt(_BODY_PT)
    style.font.color.rgb = _TITLE_COLOR
    pf = style.paragraph_format
    pf.space_after  = Pt(2)
    pf.space_before = Pt(0)


# ── Sections ────────────────────────────────────────────────────────────────

def _add_header(doc: "Document", r: dict[str, Any]) -> None:
    name = (r.get("name") or "").strip() or "Your Name"
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.font.size = Pt(_NAME_PT)
    run.font.bold = True
    run.font.color.rgb = _TITLE_COLOR
    p.paragraph_format.space_after = Pt(2)

    headline = (r.get("headline") or "").strip()
    if headline:
        p = doc.add_paragraph()
        run = p.add_run(headline)
        run.font.size = Pt(_HEADLINE_PT)
        run.italic = True
        run.font.color.rgb = _MUTED_COLOR
        p.paragraph_format.space_after = Pt(2)

    contact = r.get("contact") or {}
    bits: list[str] = []
    for key in ("email", "phone", "location", "linkedin", "github"):
        val = (contact.get(key) or "").strip()
        if val:
            bits.append(val)
    if bits:
        p = doc.add_paragraph()
        run = p.add_run(" • ".join(bits))
        run.font.size = Pt(_CONTACT_PT)
        run.font.color.rgb = _MUTED_COLOR
        p.paragraph_format.space_after = Pt(8)

    _add_divider(doc)


def _add_summary(doc: "Document", r: dict[str, Any]) -> None:
    summary = (r.get("summary") or "").strip()
    if not summary:
        return
    _add_section_heading(doc, "Professional Summary")
    p = doc.add_paragraph()
    run = p.add_run(summary)
    run.font.size = Pt(_BODY_PT)
    p.paragraph_format.space_after = Pt(4)


def _add_skills(doc: "Document", r: dict[str, Any]) -> None:
    skills = r.get("skills")
    if not skills:
        return
    _add_section_heading(doc, "Key Skills")

    if isinstance(skills, dict):
        for group, items in skills.items():
            if not items:
                continue
            p = doc.add_paragraph()
            label = p.add_run(f"{group}: ")
            label.bold = True
            label.font.size = Pt(_BODY_PT)
            body = p.add_run(" • ".join(_as_str_list(items)))
            body.font.size = Pt(_BODY_PT)
            p.paragraph_format.space_after = Pt(2)
    else:
        p = doc.add_paragraph()
        run = p.add_run(" • ".join(_as_str_list(skills)))
        run.font.size = Pt(_BODY_PT)
        p.paragraph_format.space_after = Pt(2)


def _add_experience(doc: "Document", r: dict[str, Any]) -> None:
    roles = r.get("experience") or []
    if not roles:
        return
    _add_section_heading(doc, "Professional Experience")

    for role in roles:
        if not isinstance(role, dict):
            continue
        # Title — Company, Location           Dates
        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(4)
        head.paragraph_format.space_after  = Pt(0)

        title_run = head.add_run((role.get("title") or "").strip())
        title_run.bold = True
        title_run.font.size = Pt(_BODY_PT)

        company = (role.get("company") or "").strip()
        location = (role.get("location") or "").strip()
        if company:
            sep_run = head.add_run(" — ")
            sep_run.font.size = Pt(_BODY_PT)
            company_run = head.add_run(company)
            company_run.italic = True
            company_run.font.size = Pt(_BODY_PT)
            if location:
                loc_run = head.add_run(f", {location}")
                loc_run.font.size = Pt(_BODY_PT)
                loc_run.font.color.rgb = _MUTED_COLOR

        dates = _format_dates(role.get("start"), role.get("end"))
        if dates:
            tab_run = head.add_run("\t")
            tab_run.font.size = Pt(_BODY_PT)
            date_run = head.add_run(dates)
            date_run.font.size = Pt(_DATES_PT)
            date_run.font.color.rgb = _MUTED_COLOR

        # Optional context line.
        ctx = (role.get("context") or "").strip()
        if ctx:
            cp = doc.add_paragraph()
            cr = cp.add_run(ctx)
            cr.italic = True
            cr.font.size = Pt(_DATES_PT)
            cr.font.color.rgb = _MUTED_COLOR
            cp.paragraph_format.space_after = Pt(0)

        # Bullets.
        for bullet in role.get("bullets") or []:
            text = str(bullet).strip()
            if not text:
                continue
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(1)
            run = bp.add_run(text)
            run.font.size = Pt(_BODY_PT)

        # Optional key achievement line.
        ka = (role.get("key_achievement") or "").strip()
        if ka:
            kp = doc.add_paragraph()
            label = kp.add_run("Key Achievement: ")
            label.bold = True
            label.font.size = Pt(_DATES_PT)
            body = kp.add_run(ka)
            body.font.size = Pt(_DATES_PT)
            kp.paragraph_format.space_after = Pt(4)


def _add_achievements(doc: "Document", r: dict[str, Any]) -> None:
    items = r.get("achievements") or []
    if not items:
        return
    _add_section_heading(doc, "Key Achievements")
    _add_bullet_list(doc, items)


def _add_projects(doc: "Document", r: dict[str, Any]) -> None:
    projects = r.get("projects") or []
    if not projects:
        return
    _add_section_heading(doc, "Notable Projects")

    for proj in projects:
        if not isinstance(proj, dict):
            continue
        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(2)
        head.paragraph_format.space_after  = Pt(0)
        name_run = head.add_run((proj.get("name") or "").strip())
        name_run.bold = True
        name_run.font.size = Pt(_BODY_PT)
        tech = proj.get("tech") or []
        if tech:
            t = head.add_run(f" [{', '.join(_as_str_list(tech))}]")
            t.font.size = Pt(_DATES_PT)
            t.font.color.rgb = _MUTED_COLOR

        desc = (proj.get("description") or "").strip()
        if desc:
            dp = doc.add_paragraph()
            dr = dp.add_run(desc)
            dr.font.size = Pt(_BODY_PT)
            dp.paragraph_format.space_after = Pt(2)

        for bullet in proj.get("bullets") or []:
            bp = doc.add_paragraph(style="List Bullet")
            run = bp.add_run(str(bullet).strip())
            run.font.size = Pt(_BODY_PT)


def _add_education(doc: "Document", r: dict[str, Any]) -> None:
    items = r.get("education") or []
    if not items:
        return
    _add_section_heading(doc, "Education")

    for e in items:
        if not isinstance(e, dict):
            continue
        head = doc.add_paragraph()
        head.paragraph_format.space_after = Pt(0)
        deg_run = head.add_run((e.get("degree") or "").strip())
        deg_run.bold = True
        deg_run.font.size = Pt(_BODY_PT)
        inst = (e.get("institution") or "").strip()
        if inst:
            sep = head.add_run(" — ")
            sep.font.size = Pt(_BODY_PT)
            ir = head.add_run(inst)
            ir.italic = True
            ir.font.size = Pt(_BODY_PT)
        year = (e.get("year") or "").strip()
        if year:
            tab = head.add_run("\t")
            tab.font.size = Pt(_BODY_PT)
            yr = head.add_run(year)
            yr.font.size = Pt(_DATES_PT)
            yr.font.color.rgb = _MUTED_COLOR
        details = (e.get("details") or "").strip()
        if details:
            dp = doc.add_paragraph()
            dr = dp.add_run(details)
            dr.font.size = Pt(_DATES_PT)
            dr.font.color.rgb = _MUTED_COLOR
            dp.paragraph_format.space_after = Pt(2)


def _add_certifications(doc: "Document", r: dict[str, Any]) -> None:
    items = r.get("certifications") or []
    if not items:
        return
    _add_section_heading(doc, "Certifications & Training")

    for cert in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(cert, dict):
            name = (cert.get("name") or "").strip()
            issuer = (cert.get("issuer") or "").strip()
            year = (cert.get("year") or "").strip()
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(_BODY_PT)
            tail = ""
            if issuer:
                tail += f" — {issuer}"
            if year:
                tail += f" ({year})"
            if tail:
                t = p.add_run(tail)
                t.font.size = Pt(_BODY_PT)
        else:
            run = p.add_run(str(cert).strip())
            run.font.size = Pt(_BODY_PT)


def _add_publications(doc: "Document", r: dict[str, Any]) -> None:
    items = r.get("publications") or []
    if not items:
        return
    _add_section_heading(doc, "Publications & Patents")
    _add_bullet_list(doc, items)


def _add_awards(doc: "Document", r: dict[str, Any]) -> None:
    items = r.get("awards") or []
    if not items:
        return
    _add_section_heading(doc, "Awards & Recognition")
    _add_bullet_list(doc, items)


def _add_memberships(doc: "Document", r: dict[str, Any]) -> None:
    items = r.get("memberships") or []
    if not items:
        return
    _add_section_heading(doc, "Professional Memberships")
    p = doc.add_paragraph()
    run = p.add_run(" • ".join(_as_str_list(items)))
    run.font.size = Pt(_BODY_PT)


def _add_languages(doc: "Document", r: dict[str, Any]) -> None:
    items = r.get("languages") or []
    if not items:
        return
    _add_section_heading(doc, "Languages")
    p = doc.add_paragraph()
    run = p.add_run(" • ".join(_as_str_list(items)))
    run.font.size = Pt(_BODY_PT)


# ── Primitives ──────────────────────────────────────────────────────────────

def _add_section_heading(doc: "Document", text: str) -> None:
    """A subtle, uppercase heading with a divider underneath."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(_SECTION_PT)
    run.font.color.rgb = _TITLE_COLOR
    _add_divider(doc)


def _add_divider(doc: "Document") -> None:
    """A thin grey rule made of an underlined paragraph of spaces.

    We avoid the XML-bottom-border trick to keep this dependency-free; a
    simple long underscore does the job and edits cleanly in Word.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("_" * 95)
    run.font.size = Pt(6)
    run.font.color.rgb = _DIVIDER_COLOR


def _add_bullet_list(doc: "Document", items: Iterable[Any]) -> None:
    for item in items:
        text = str(item).strip()
        if not text:
            continue
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(_BODY_PT)


# ─────────────────────────────────────────────────────────────────────────────
# Interview Prep DOCX
# ─────────────────────────────────────────────────────────────────────────────

def render_interview_prep_docx(
    prep_data: dict[str, Any],
    template: str = "executive",
) -> bytes:
    """Render an interview prep kit as a clean, editable Word document."""
    if template != "executive":
        pass  # only one template today; fall through.

    doc = Document()
    _configure_document(doc)

    p = prep_data or {}
    ctx = p.get("context") or {}

    _add_prep_header(doc, ctx)
    _add_prep_company_brief(doc, p.get("company_brief") or {})
    _add_prep_likely_questions(doc, p.get("likely_questions") or [])
    _add_prep_suggested_answers(doc, p.get("suggested_answers") or [])
    _add_prep_questions_to_ask(doc, p.get("questions_to_ask") or [])
    _add_prep_red_flags(doc, p.get("red_flags") or [])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_prep_header(doc: "Document", ctx: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    run = p.add_run("Interview Prep Kit")
    run.font.size = Pt(_NAME_PT)
    run.font.bold = True
    run.font.color.rgb = _TITLE_COLOR
    p.paragraph_format.space_after = Pt(2)

    bits: list[str] = []
    if ctx.get("target_role"): bits.append(str(ctx["target_role"]))
    if ctx.get("company"):     bits.append(str(ctx["company"]))
    if bits:
        p = doc.add_paragraph()
        run = p.add_run(" — ".join(bits))
        run.font.size = Pt(_HEADLINE_PT)
        run.italic = True
        run.font.color.rgb = _MUTED_COLOR

    if ctx.get("round_label"):
        p = doc.add_paragraph()
        run = p.add_run(f"Round: {ctx['round_label']}")
        run.font.size = Pt(_CONTACT_PT)
        run.font.color.rgb = _MUTED_COLOR

    _add_divider(doc)


def _add_prep_company_brief(doc: "Document", cb: dict[str, Any]) -> None:
    if not cb:
        return
    _add_section_heading(doc, "Company Brief")

    if cb.get("summary"):
        p = doc.add_paragraph()
        run = p.add_run(str(cb["summary"]))
        run.font.size = Pt(_BODY_PT)

    def _sub(label: str, items: Any) -> None:
        items = items or []
        if not items:
            return
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(4)
        sp.paragraph_format.space_after  = Pt(2)
        run = sp.add_run(label)
        run.bold = True
        run.font.size = Pt(_DATES_PT)
        run.font.color.rgb = _TITLE_COLOR
        _add_bullet_list(doc, items)

    _sub("Recent News",            cb.get("recent_news"))
    _sub("Culture & Working Style", cb.get("culture"))

    if cb.get("investor_thesis"):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(4)
        run = sp.add_run("Investor Thesis")
        run.bold = True
        run.font.size = Pt(_DATES_PT)
        run.font.color.rgb = _TITLE_COLOR
        bp = doc.add_paragraph()
        br = bp.add_run(str(cb["investor_thesis"]))
        br.font.size = Pt(_BODY_PT)

    _sub("Strategic Challenges", cb.get("challenges"))


def _add_prep_likely_questions(
    doc: "Document", qs: list[dict[str, Any]],
) -> None:
    if not qs:
        return
    _add_section_heading(doc, "Likely Questions")
    for i, q in enumerate(qs, start=1):
        if not isinstance(q, dict):
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(1)
        num = p.add_run(f"Q{i}. ")
        num.bold = True
        num.font.size = Pt(_BODY_PT)
        body = p.add_run(str(q.get("question") or ""))
        body.font.size = Pt(_BODY_PT)
        if q.get("category"):
            tag = p.add_run(f"  [{q['category']}]")
            tag.font.size = Pt(_DATES_PT)
            tag.italic = True
            tag.font.color.rgb = _MUTED_COLOR
        if q.get("why_asked"):
            wp = doc.add_paragraph()
            wp.paragraph_format.left_indent = Cm(0.6)
            wp.paragraph_format.space_after = Pt(2)
            wr = wp.add_run(f"Why: {q['why_asked']}")
            wr.italic = True
            wr.font.size = Pt(_DATES_PT)
            wr.font.color.rgb = _MUTED_COLOR


def _add_prep_suggested_answers(
    doc: "Document", answers: list[dict[str, Any]],
) -> None:
    if not answers:
        return
    _add_section_heading(doc, "Suggested Answers (STAR)")
    for sa in answers:
        if not isinstance(sa, dict):
            continue
        # Question line
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(4)
        qp.paragraph_format.space_after  = Pt(1)
        qr = qp.add_run(str(sa.get("question") or ""))
        qr.bold = True
        qr.font.size = Pt(_BODY_PT)

        star = sa.get("star") or {}
        for label, key in (("Situation", "situation"), ("Task", "task"),
                           ("Action", "action"), ("Result", "result")):
            val = (star.get(key) or "").strip()
            if not val:
                continue
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(1)
            l = sp.add_run(f"{label}: ")
            l.bold = True
            l.font.size = Pt(_DATES_PT)
            l.font.color.rgb = _TITLE_COLOR
            b = sp.add_run(val)
            b.font.size = Pt(_BODY_PT)

        if sa.get("anchor_evidence"):
            ap = doc.add_paragraph()
            ap.paragraph_format.space_after = Pt(2)
            ar = ap.add_run(f"Anchored in: {sa['anchor_evidence']}")
            ar.italic = True
            ar.font.size = Pt(_DATES_PT)
            ar.font.color.rgb = _MUTED_COLOR


def _add_prep_questions_to_ask(
    doc: "Document", qs: list[dict[str, Any]],
) -> None:
    if not qs:
        return
    _add_section_heading(doc, "Questions to Ask")
    for q in qs:
        if not isinstance(q, dict):
            continue
        bp = doc.add_paragraph(style="List Bullet")
        run = bp.add_run(str(q.get("question") or ""))
        run.font.size = Pt(_BODY_PT)
        run.bold = True
        if q.get("signals"):
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Cm(0.6)
            sp.paragraph_format.space_after = Pt(2)
            sr = sp.add_run(f"Signals: {q['signals']}")
            sr.italic = True
            sr.font.size = Pt(_DATES_PT)
            sr.font.color.rgb = _MUTED_COLOR


def _add_prep_red_flags(
    doc: "Document", reds: list[dict[str, Any]],
) -> None:
    if not reds:
        return
    _add_section_heading(doc, "Red Flags — Avoid These")
    for r in reds:
        if not isinstance(r, dict):
            continue
        mp = doc.add_paragraph()
        mp.paragraph_format.space_before = Pt(3)
        mp.paragraph_format.space_after  = Pt(1)
        l = mp.add_run("Mistake: ")
        l.bold = True
        l.font.size = Pt(_DATES_PT)
        b = mp.add_run(str(r.get("mistake") or ""))
        b.font.size = Pt(_BODY_PT)

        if r.get("instead"):
            ip = doc.add_paragraph()
            ip.paragraph_format.left_indent = Cm(0.4)
            ip.paragraph_format.space_after = Pt(2)
            il = ip.add_run("Instead: ")
            il.bold = True
            il.font.size = Pt(_DATES_PT)
            ib = ip.add_run(str(r["instead"]))
            ib.font.size = Pt(_BODY_PT)


def _format_dates(start: Any, end: Any) -> str:
    s = (str(start).strip() if start else "")
    e = (str(end).strip() if end else "")
    if s and e:
        return f"{s} – {e}"
    return s or e


def _as_str_list(values: Any) -> list[str]:
    if not values:
        return []
    return [str(v).strip() for v in values if str(v).strip()]
